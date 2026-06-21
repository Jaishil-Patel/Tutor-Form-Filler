"""Build the Postgraduate Merit Award Duties Timesheet .docx to exact spec.

Each spec section from the form specification is implemented as a small helper.
The only third-party dependency is python-docx.
"""

from __future__ import annotations

import os
from typing import List

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

from . import config
from .models import Session

RED = RGBColor(0xFF, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)
FONT_NAME = "Calibri"
BODY_PT = 11

# Main table column widths (cm), proportional to the spec.
MAIN_COL_WIDTHS_CM = [2.3, 3.3, 2.0, 2.3, 1.8, 2.4, 2.4]
MAIN_HEADERS = [
    "Date of work",
    "Duties/Description of work done",
    "Time Started",
    "Time Ended",
    "Total Hours",
    "Student signature",
    "Supervisor Signature",
]

SUBMISSION_COL_WIDTHS_CM = [2.2, 4.0, 2.5]


# --------------------------------------------------------------------------
# Low-level helpers
# --------------------------------------------------------------------------
def _set_run_font(run, *, size=BODY_PT, bold=False, italic=False,
                  underline=False, color=BLACK, superscript=False):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.color.rgb = color
    run.font.superscript = superscript
    return run


def _blank_paragraph(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    return p


def _set_cell_text(cell, lines, *, bold=False, color=BLACK):
    """Replace a cell's content with one or more lines of styled text."""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    if isinstance(lines, str):
        lines = [lines]
    para = cell.paragraphs[0]
    para.text = ""
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    first = True
    for line in lines:
        p = para if first else cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line)
        _set_run_font(run, bold=bold, color=color)
        first = False


def _set_table_borders(table):
    """Apply thin single black borders to every cell of the table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")  # 4 eighths of a point = thin
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tblPr.append(borders)


def _fix_column_widths(table, widths_cm):
    """Force fixed column widths (python-docx needs width set on every cell)."""
    table.autofit = False
    table.allow_autofit = False
    # mark layout as fixed
    tblPr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)


# --------------------------------------------------------------------------
# Section builders
# --------------------------------------------------------------------------
def _setup_page(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(BODY_PT)
    normal.font.color.rgb = BLACK


def _add_logo(doc, settings):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_path = config.resolve_logo_path(settings)
    width_cm = float(settings.get("logo_width_cm", 4.3))
    if logo_path and os.path.exists(logo_path):
        run = p.add_run()
        run.add_picture(logo_path, width=Cm(width_cm))
    else:
        run = p.add_run("[University of the Witwatersrand logo]")
        _set_run_font(run, italic=True)
    _blank_paragraph(doc)


def _add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("POSTGRADUATE MERIT AWARD DUTIES TIMESHEET")
    _set_run_font(run, bold=True, italic=True, underline=True)
    _blank_paragraph(doc)


def _add_student_info(doc, settings, month_label):
    rows = [
        ("Student Name: ", settings.get("student_name", "")),
        ("Student No: ", settings.get("student_no", "")),
        ("School: ", settings.get("school", "")),
        ("Month of Claim: ", month_label),
    ]
    for label, value in rows:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(f"{label}{value}")
        _set_run_font(run, bold=True)
    _blank_paragraph(doc)


def _add_main_table(doc, sessions: List[Session]):
    n_rows = 1 + len(sessions)
    table = doc.add_table(rows=n_rows, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table)

    # header row
    for idx, heading in enumerate(MAIN_HEADERS):
        _set_cell_text(table.rows[0].cells[idx], heading, bold=True, color=BLACK)

    # data rows
    for r, s in enumerate(sessions, start=1):
        cells = table.rows[r].cells
        _set_cell_text(cells[0], s.display_date(), bold=False, color=RED)
        duties = [s.course_code]
        if s.activity:
            duties.append(s.activity)
        _set_cell_text(cells[1], duties, bold=False, color=BLACK)
        _set_cell_text(cells[2], s.time_started, bold=False, color=BLACK)
        _set_cell_text(cells[3], s.time_ended, bold=False, color=BLACK)
        _set_cell_text(cells[4], s.effective_hours(), bold=False, color=BLACK)
        _set_cell_text(cells[5], "", bold=False, color=BLACK)
        _set_cell_text(cells[6], "", bold=False, color=BLACK)

    _fix_column_widths(table, MAIN_COL_WIDTHS_CM)

    # give rows enough height for two text lines + a signature
    for row in table.rows[1:]:
        row.height = Cm(1.1)

    _blank_paragraph(doc)


def _add_hos_line(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    label = p.add_run("HOS/HOD Signature: ")
    _set_run_font(label, bold=True)
    rule = p.add_run("______________________")
    _set_run_font(rule, bold=False)
    _blank_paragraph(doc)
    _blank_paragraph(doc)


def _add_submission_intro(doc, settings):
    intro = settings.get("submission_intro", "")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Render with only "due date" bold + underlined.
    marker = "due date"
    if marker in intro:
        before, after = intro.split(marker, 1)
        _set_run_font(p.add_run(before))
        _set_run_font(p.add_run(marker), bold=True, underline=True)
        _set_run_font(p.add_run(after))
    else:
        _set_run_font(p.add_run(intro))
    _blank_paragraph(doc)


def _add_submission_table(doc, settings):
    rows = settings.get("submission_rows", [])
    table = doc.add_table(rows=len(rows), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_borders(table)

    for r, row in enumerate(rows):
        cells = table.rows[r].cells
        for c in cells:
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        # Column 1: block label (bold) with superscript ordinal suffix
        p0 = cells[0].paragraphs[0]
        p0.text = ""
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_run_font(p0.add_run(row.get("block_num", "")), bold=True)
        _set_run_font(p0.add_run(row.get("block_suffix", "")), bold=True,
                      superscript=True)
        _set_run_font(p0.add_run(" Block"), bold=True)

        # Column 2: date with superscript ordinal on the day number
        p1 = cells[1].paragraphs[0]
        p1.text = ""
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_run_font(p1.add_run(row.get("date_day", "")))
        _set_run_font(p1.add_run(row.get("date_suffix", "")), superscript=True)
        _set_run_font(p1.add_run(row.get("date_rest", "")))

        # Column 3: hours
        _set_cell_text(cells[2], row.get("hours", ""))

    _fix_column_widths(table, SUBMISSION_COL_WIDTHS_CM)
    _blank_paragraph(doc)


def _add_bullets(doc, settings):
    for bullet in settings.get("bullets", []):
        text = bullet.get("text", "")
        prefix = bullet.get("underline_prefix", "")
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if prefix and text.startswith(prefix):
            _set_run_font(p.add_run(prefix), bold=True, underline=True, color=RED)
            _set_run_font(p.add_run(text[len(prefix):]), bold=True, color=RED)
        else:
            _set_run_font(p.add_run(text), bold=True, color=RED)


# --------------------------------------------------------------------------
# Public entry point
# --------------------------------------------------------------------------
def build_timesheet(sessions: List[Session], settings: dict, month_label: str,
                    out_path: str) -> str:
    """Build the timesheet .docx and write it to out_path. Returns out_path."""
    doc = Document()
    _setup_page(doc)
    _add_logo(doc, settings)
    _add_title(doc)
    _add_student_info(doc, settings, month_label)
    _add_main_table(doc, sessions)
    _add_hos_line(doc)
    _add_submission_intro(doc, settings)
    _add_submission_table(doc, settings)
    _add_bullets(doc, settings)

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    return out_path
